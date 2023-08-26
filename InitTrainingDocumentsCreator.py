from docx import Document
from random import randint, choice
import datetime
import sys


SOURCE_FILES = "Data/sources"
INPUT_DOCUMENT = f"{SOURCE_FILES}/hospitalInformationSheet.docx"
OUTPUT_DIRECTORY = "Data/generated_documents"


def err_exit(*args):
    print('Error: '*args, file=sys.stderr)
    exit(1)


class HospitalInformationSheetCreator:
    name = str()
    sex = str()
    pesel = str()
    address = str()
    phone_number = str()
    birthdate = None
    arrival_date = None
    discharge_date = None

    def _create_name(self):
        if self.sex == 'M':
            file_first_names = open(f'{SOURCE_FILES}/male_firstnames.csv', 'r', encoding='utf-8')
            file_last_names = open(f'{SOURCE_FILES}/male_lastnames.csv', 'r', encoding='utf-8')
        else:
            file_first_names = open(f'{SOURCE_FILES}/female_firstnames.csv', 'r', encoding='utf-8')
            file_last_names = open(f'{SOURCE_FILES}/female_lastnames.csv', 'r', encoding='utf-8')

        first_names = file_first_names.readlines()
        last_names = file_last_names.readlines()

        file_first_names.close()
        file_last_names.close()

        self.name = choice(first_names).rstrip('\n').capitalize() + ' ' + choice(last_names).rstrip('\n').capitalize()

    def _create_dates(self):
        """
        Returns list of length 5:
        [0] - birthdate
        [1] - arrival at hospital date
        [2] - arrival at hospital hour
        [3] - discharge date
        [4] - discharge hour
        """

        # generate birthdate date from 23 to 60 years ago
        self.birthdate = datetime.date.today() - datetime.timedelta(days=(randint(23, 60) * 365))

        # Generate a random arrival date between 7 days and 20 years ago
        self.arrival_date = datetime.datetime.now() - datetime.timedelta(days=randint(7, 365 * 20))
        self.arrival_date = self.arrival_date.replace(hour=randint(0, 23), minute=randint(0, 59))

        # Generate random discharge date between 1 and 7 days after the arrival date
        self.discharge_date = self.arrival_date + datetime.timedelta(days=randint(1, 5))
        self.discharge_date = self.discharge_date.replace(hour=randint(10, 16), minute=choice([0, 30]))

        return [birthdate, arrival_date.strftime('%Y-%m-%d'), arrival_date.strftime('%H:%M'),
                discharge_date.strftime('%Y-%m-%d'), discharge_date.strftime('%H:%M')]

    def _create_address(self):
        addresses_file = open(f'{SOURCE_FILES}/addresses.csv', 'r', encoding='utf-8')
        addresses = addresses_file.readlines()
        addresses_file.close()

        self.address = choice(addresses).rstrip('\n') + ' ' + str(randint(1, 100))
        if randint(1, 100) > 55:  # chance for someone to live in a flat
            self.address += '/' + randint(1, 80)

    def _create_phone_number(self):
        self.phone_number = str(randint(5, 8))  # first digit
        self.phone_number += ''.join([str(randint(0, 9)) for _ in range(8)])  # remaining digits

    def _create_pesel(self):
        if self.birthdate is None:
            err_exit('Birthdate is not set when _create_pesel was called')

        p_year = str(self.birthdate.year)[-2:]  # last two digits of year

        p_month = str()
        if 1800 <= self.birthdate.year < 1900:
            p_month = str(self.birthdate.month + 80)
        elif self.birthdate.year < 2000:
            p_month = str(self.birthdate.month)
        elif self.birthdate.year < 2100:
            p_month = str(self.birthdate.month + 20)
        else:
            err_exit('Birth year is not between 1800-2100')

        p_day = str(self.birthdate.day)

        p_ordinal_number = ''.join([str(randint(0, 9)) for _ in range(3)])
        if self.sex == 'M':
            p_ordinal_number += str(randint(0, 4) * 2 + 1)
        elif self.sex == 'F':
            p_ordinal_number += str(randint(0, 4) * 2)
        else:
            err_exit('Unknown sex')

        pesel_tmp = p_year + p_month + p_day + p_ordinal_number

        factors = [1, 3, 7, 9]
        result_sum = 0
        for i in range(10):
            digit = pesel_tmp[i]
            result_sum += digit * factors[i % 4]
        p_control_number = result_sum % 10

        self.pesel = pesel_tmp + p_control_number

    def produce_new_file(self):
        information_sheet = Document(INPUT_DOCUMENT)
        self.sex = choice(['K', 'M'])
        self._create_pesel()
        self._create_name()
        self._create_dates()
        self._create_address()
        self._create_phone_number()

        paragraph = information_sheet.paragraphs[3]
        paragraph.runs[6].text = f'{self.name}    '
        paragraph.runs[11].text = f'{self.pesel}'

        paragraph = information_sheet.paragraphs[4]
        paragraph.runs[3].text = f"Kraków, {self.discharge_date.strftime('%Y-%m-%d')}"
        # TODO: dodać miasta do adresów

        paragraph = information_sheet.paragraphs[10]
        paragraph.runs[2].text = f'{new_name}'
        paragraph.runs[7].text = f'{new_pesel}'

        paragraph = information_sheet.paragraphs[11]
        paragraph.runs[3].text = f'{new_dates[0]}'
        paragraph.runs[11].text = f'{new_sex}'

        paragraph = information_sheet.paragraphs[12]
        paragraph.runs[0].text = f'Adres zamieszkania: {new_address}'

        paragraph = information_sheet.paragraphs[13]
        paragraph.runs[0].text = f'Data i godzina przyjęcia pacjenta: {new_dates[1]}, {new_dates[2]}, ' \
                                 f'data i godzina wypisu pacjenta: {new_dates[3]}, {new_dates[4]}'

        information_sheet.save(f'{OUTPUT_DIRECTORY}/edited.docx')



def create_xml(information_sheet):
    with open('HospitalInformationSheet.xml', 'w') as f:
        f.write(information_sheet._element.xml)


def print_docx(information_sheet):
    i = 0
    for paragraph in information_sheet.paragraphs:
        print(f'{i}: {paragraph.text}')
        i += 1


