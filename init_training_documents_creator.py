import glob
import os
from docx import Document

import fake_data_creator
import utils


SOURCE_FILES = "data/sources"
INPUT_DOCUMENT = f"{SOURCE_FILES}/hospitalInformationSheet.docx"
OUTPUT_DIRECTORY = "output"


def produce_new_init_document(filename):
    """
    Creates and saves a new hospital information sheet as docx document with randomized data
    :param filename: how should the output file be named
    :return: Returns 0 if successful, 1 if initial file doesn't exist
    """

    if not os.path.exists(INPUT_DOCUMENT):
        print("Error: Could not find initial training document HospitalInformationSheet.docx")
        return 1
    information_sheet = Document(INPUT_DOCUMENT)

    fake_data = fake_data_creator.FakeDataCreator()
    person = fake_data.create_person()
    address = fake_data.create_address()
    dates = fake_data.create_dates()
    phone_number = fake_data.create_phone_number()

    birthday = person['birthdate']
    start_date = dates['start_date']
    end_date = dates['end_date']

    paragraph = information_sheet.paragraphs[3]
    paragraph.runs[6].text = f"{person['first_name']} {person['last_name']}"
    paragraph.runs[11].text = f"{person['pesel']}"

    paragraph = information_sheet.paragraphs[4]
    paragraph.runs[3].text = f"Kraków, {end_date.strftime('%Y-%m-%d')}"

    paragraph = information_sheet.paragraphs[10]
    paragraph.runs[2].text = f"{person['first_name']} {person['last_name']}"
    paragraph.runs[7].text = f"{person['pesel']}"

    paragraph = information_sheet.paragraphs[11]
    paragraph.runs[3].text = f"{birthday.strftime('%Y-%m-%d')}"
    paragraph.runs[6].text = f" {phone_number}"
    paragraph.runs[13].text = f"{person['sex']}"

    paragraph = information_sheet.paragraphs[12]
    paragraph.runs[0].text = f"Adres zamieszkania: {address['street']}, {address['city']}"

    paragraph = information_sheet.paragraphs[13]
    paragraph.runs[0].text = f"Data i godzina przyjęcia pacjenta: {start_date.strftime('%Y-%m-%d')}, " \
                             f"{start_date.strftime('%H:%M')}, " \
                             f"data i godzina wypisu pacjenta: {end_date.strftime('%y-%m-%d')}, " \
                             f"{end_date.strftime('%H:%M')}"

    information_sheet.save(f"{OUTPUT_DIRECTORY}/{filename}")

    return 0


def training_docx_to_bmp():
    """
    Converts training images to jpg
    :return: Returns 0 if successful, 1 if it couldn't find any docx at Data/sources/training_documents
    """

    # change docx to pdf
    docx_files = glob.glob(os.path.join(OUTPUT_DIRECTORY, "*.docx"))
    if not docx_files:
        print("Warning: not docx files found to convert to bmp!")
        return 1
    for word_file in docx_files:
        utils.docx_to_bmp(word_file)
        os.remove(word_file)

    return 0


def generate_training_documents(amount):
    if not os.path.exists(OUTPUT_DIRECTORY):
        os.makedirs(OUTPUT_DIRECTORY)

    for i in range(amount):
        output_name = f"training_document_{i}.docx"
        if produce_new_init_document(output_name):
            return 1

    if training_docx_to_bmp():
        return 1
